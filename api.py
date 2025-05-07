from fastapi import FastAPI, Request, File, UploadFile, Form, Body
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse
import shutil
from pathlib import Path
import json
import numpy as np
import ollama
import requests
import openai
from transformers import pipeline

from pydantic import BaseModel
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import Qdrant
from langchain_community.llms import Ollama
from langchain_ollama import OllamaEmbeddings, OllamaLLM
from langchain.chains import ConversationalRetrievalChain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.http.models import VectorParams, Distance
import os
from datetime import datetime
import pandas as pd  # Add this import for Excel processing
from docx import Document  # For RFP doc generation
from fpdf import FPDF  # For PDF generation from proposal text
import re
from fastapi.openapi.utils import get_openapi

app = FastAPI()

class ChatRequest(BaseModel):
    question: str
    chat_history: list = []
    rfp_service: bool = False  # Add this field
    rfp_info: dict = None      # Add this field for user RFP info

class ProposalUpdateRequest(BaseModel):
    proposal_text: str
    update_instructions: str
    chat_history: list = []

# Define paths for original and uploaded documents
DOCS_PATH = os.path.join(os.path.dirname(__file__), 'doccuments')
UPLOADS_PATH = os.path.join(os.path.dirname(__file__), 'uploads')

# Create uploads folder if it doesn't exist
if not os.path.exists(UPLOADS_PATH):
    os.makedirs(UPLOADS_PATH)

# Define separate collection names
ORIGINAL_COLLECTION_NAME = "bangla_chatbot_docs"
UPLOADS_COLLECTION_NAME = "bangla_chatbot_uploads"
FRAMEWORK_AGREEMENTS_COLLECTION = "framework_agreements"

# Setup for original documents
pdf_files = [os.path.join(DOCS_PATH, f) for f in os.listdir(DOCS_PATH) if f.endswith('.pdf')]
documents = []
for pdf in pdf_files:
    loader = PyPDFLoader(pdf)
    documents.extend(loader.load())
splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
chunks = splitter.split_documents(documents)
model = SentenceTransformer('all-MiniLM-L6-v2')
embeddings = [model.encode(chunk.page_content) for chunk in chunks]
QDRANT_HOST = os.getenv("QDRANT_HOST", "localhost")
QDRANT_PORT = int(os.getenv("QDRANT_PORT", "6333"))

# Setup Qdrant client
qdrant_client = QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT)

# Initialize original documents collection
try:
    qdrant_client.delete_collection(collection_name=ORIGINAL_COLLECTION_NAME)
except Exception as e:
    print(f"Original collection delete skipped or failed: {e}")
try:
    qdrant_client.create_collection(
        collection_name=ORIGINAL_COLLECTION_NAME,
        vectors_config=VectorParams(size=384, distance=Distance.COSINE),
    )
except Exception as e:
    print(f"Original collection creation skipped or failed: {e}")

# Initialize uploads collection if it doesn't exist
try:
    if not qdrant_client.collection_exists(UPLOADS_COLLECTION_NAME):
        qdrant_client.create_collection(
            collection_name=UPLOADS_COLLECTION_NAME,
            vectors_config=VectorParams(size=384, distance=Distance.COSINE),
        )
except Exception as e:
    print(f"Uploads collection creation skipped or failed: {e}")

# Initialize framework agreements collection if it doesn't exist
try:
    if not qdrant_client.collection_exists(FRAMEWORK_AGREEMENTS_COLLECTION):
        qdrant_client.create_collection(
            collection_name=FRAMEWORK_AGREEMENTS_COLLECTION,
            vectors_config=VectorParams(size=384, distance=Distance.COSINE),
        )
except Exception as e:
    print(f"Framework agreements collection creation skipped or failed: {e}")

# Insert original documents into their collection
payloads = [{"text": chunk.page_content, "source": "original"} for chunk in chunks]
if payloads:  # Only insert if there are documents
    qdrant_client.upsert(
        collection_name=ORIGINAL_COLLECTION_NAME,
        points=[
            {
                "id": i,
                "vector": embeddings[i],
                "payload": payloads[i],
            }
            for i in range(len(embeddings))
        ],
    )

# Remove Ollama and use Hugging Face pipeline for LLM
llm_pipeline = pipeline("text-generation", model="gpt2")  # You can change to any open LLM
GEMMA_API_URL = "http://localhost:8000/v1/chat/completions"  # Update if your endpoint is different
GEMMA_API_KEY = os.getenv("GEMMA_API_KEY", "")  # Set in your environment if needed

OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "http://localhost:12434/engines/v1")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "ignored")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "ai/gemma3")

openai_client = openai.OpenAI(
    base_url=OPENAI_BASE_URL,
    api_key=OPENAI_API_KEY,
)

@app.post("/chat")
def chat_endpoint(request: ChatRequest):
    def extract_proposal(text):
        match = re.search(r'<RFP_PROPOSAL>(.*?)</RFP_PROPOSAL>', text, re.DOTALL)
        return match.group(1).strip() if match else None

    def generate_pdf_from_text(text, filename):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for line in text.splitlines():
            pdf.multi_cell(0, 10, line)
        pdf.output(filename)

    def stream_response():
        # If RFP generation is requested
        if getattr(request, 'rfp_service', False):
            # 1. Retrieve BRAC guideline context from vector store
            guideline_query = "BRAC RFP guideline"
            guideline_vec = model.encode(guideline_query)
            guideline_hits = qdrant_client.search(
                collection_name=ORIGINAL_COLLECTION_NAME,
                query_vector=guideline_vec,
                limit=3,
            )
            guideline_text = "\n".join([hit.payload['text'] for hit in guideline_hits])
            # 2. Compose RFP document using user info and guideline
            doc = Document()
            doc.add_heading('Request for Proposal (RFP)', 0)
            if request.rfp_info:
                for k, v in request.rfp_info.items():
                    doc.add_paragraph(f"{k}: {v}")
            doc.add_heading('BRAC RFP Guidelines', level=1)
            doc.add_paragraph(guideline_text)
            # 3. Save RFP file
            rfp_dir = os.path.join(os.path.dirname(__file__), 'static')
            if not os.path.exists(rfp_dir):
                os.makedirs(rfp_dir)
            rfp_filename = f"rfp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            rfp_path = os.path.join(rfp_dir, rfp_filename)
            doc.save(rfp_path)
            # 4. Yield a chat message with download link
            download_url = f"/static/{rfp_filename}"
            msg = f"<h2>RFP Generated</h2><p>Your RFP has been generated based on your input and BRAC guidelines.</p><a href=\"{download_url}\" target=\"_blank\"><strong>Download RFP Document</strong></a>"
            yield msg.encode('utf-8')
            return
        # Embed the question for similarity search
        question_vec = model.encode(request.question)
        
        # Store the current question in chat context
        chat_context.add_question(request.question, question_vec)
        
        # Get similar previous questions
        similar_questions = chat_context.get_similar_questions(question_vec)
        
        # Search both Qdrant collections for document context
        original_hits = qdrant_client.search(
            collection_name=ORIGINAL_COLLECTION_NAME,
            query_vector=question_vec,
            limit=5,
        )
        
        upload_hits = qdrant_client.search(
            collection_name=UPLOADS_COLLECTION_NAME,
            query_vector=question_vec,
            limit=10,
        )
        
        # Format document context with source indications
        document_context = ""
        
        if original_hits:
            document_context += "\nFROM BASE DOCUMENTS:\n"
            for hit in original_hits:
                document_context += f"Source: Original documentation\nContent: {hit.payload['text']}\n\n"
        
        if upload_hits:
            document_context += "\nFROM RECENTLY UPLOADED DOCUMENTS:\n"
            for hit in upload_hits:
                document_context += f"Source: {hit.payload['source']}\nContent: {hit.payload['text']}\n\n"
        
        # Get file contexts
        file_contexts = chat_context.get_file_contexts()
        file_context_str = ""
        if file_contexts:
            file_context_str = "\nRECENTLY UPLOADED FILES SUMMARY:\n"
            for ctx in file_contexts[-3:]:  # Use only the 3 most recent files
                file_context_str += f"Filename: {ctx.get('filename')}\nSummary: {ctx.get('summary')}\n\n"
        
        # Build previous conversation context
        conversation_context = ""
        if similar_questions:
            conversation_context = "\nRELATED PREVIOUS QUESTIONS:\n"
            for q, sim in similar_questions:
                conversation_context += f"- {q} (similarity: {sim:.2f})\n"
        
        # Combine all context
        combined_context = document_context + file_context_str + conversation_context
        
        # Prepare messages for API call
        messages = [
            {"role": "system", "content": """You are BRAC's Procurement Assistant named BRACGPT. You are an experienced procurement professional working for BRAC.

INSTRUCTIONS:
1. Respond as a knowledgeable procurement professional working for BRAC.
2. Format your responses in a clean, structured manner using HTML tags when appropriate.
3. Use <h2> for section headings, <ul> or <ol> for lists, <strong> for emphasis, and <p> for paragraphs.
4. When discussing financial data, use <table> tags with proper formatting.
5. If referencing specific sections from procurement documents, clearly cite the reference.
6. If you cannot provide an answer, politely state that you'll need to consult with the procurement team.
7. Be concise but thorough in your explanations.
8. If a user's question is similar to a previous question, acknowledge this and build upon previous responses.
9. Never reveal that you're using any default documents or stored knowledge base.
10. For uploaded documents, simply respond as if you've personally reviewed them as a procurement officer.
11. If generating an RFP proposal, enclose the proposal text within <RFP_PROPOSAL>...</RFP_PROPOSAL> tags.

Your goal is to provide helpful procurement guidance while maintaining a professional BRAC organizational identity."""},
            {"role": "user", "content": f"Context: {combined_context}\n\nQuestion: {request.question}"}
        ]
        
        data = {
            "model": OPENAI_MODEL,
            "messages": messages,
            "max_tokens": 25600,
            "temperature": 0.5,
            "n": 1,
            "stream": True
        }
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {OPENAI_API_KEY}"
        }
        with requests.post(OPENAI_BASE_URL + "/chat/completions", headers=headers, json=data, stream=True) as resp:
            full_response = ""
            for line in resp.iter_lines():
                if line:
                    if line.startswith(b'data: '):
                        line = line[6:]
                    if line == b'[DONE]':
                        break
                    try:
                        chunk = json.loads(line)
                        content = chunk.get('choices', [{}])[0].get('delta', {}).get('content', '')
                        if content:
                            full_response += content
                            yield content.encode('utf-8')  # Yield bytes for StreamingResponse
                    except Exception:
                        continue
            # After streaming, check for RFP proposal tag and generate PDF if present
            proposal = extract_proposal(full_response)
            if proposal:
                pdf_dir = os.path.join(os.path.dirname(__file__), 'static')
                if not os.path.exists(pdf_dir):
                    os.makedirs(pdf_dir)
                pdf_filename = f"rfp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                pdf_path = os.path.join(pdf_dir, pdf_filename)
                generate_pdf_from_text(proposal, pdf_path)
                download_url = f"/static/{pdf_filename}"
                msg = f"<h2>RFP Proposal PDF</h2><p>Your RFP proposal has been generated as a PDF.</p><a href=\"{download_url}\" target=\"_blank\"><strong>Download RFP PDF</strong></a>"
                yield msg.encode('utf-8')
    return StreamingResponse(stream_response(), media_type="text/plain")

@app.post("/proposal-update")
def proposal_update_endpoint(request: ProposalUpdateRequest = Body(...)):
    def generate_pdf_from_text(text, filename):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for line in text.splitlines():
            pdf.multi_cell(0, 10, line)
        pdf.output(filename)

    def stream_response():
        # Compose context for the LLM
        messages = [
            {"role": "system", "content": '''You are a senior procurement proposal editor for BRAC. Your job is to update, revise, or modify procurement proposals based on user instructions. 

INSTRUCTIONS:
1. Carefully read the user's update instructions and the current proposal text.
2. Make only the requested changes, keeping the rest of the proposal intact.
3. Respond with the full, updated proposal, formatted in HTML for clarity.
4. If the user requests a summary of changes, provide a bullet list before the updated proposal.
5. If the user requests a PDF, enclose the proposal in <RFP_PROPOSAL>...</RFP_PROPOSAL> tags so a PDF can be generated.
6. Be precise, professional, and concise.'''
            },
            {"role": "user", "content": f"Current Proposal:\n{request.proposal_text}\n\nUpdate Instructions:\n{request.update_instructions}"}
        ]
        data = {
            "model": OPENAI_MODEL,
            "messages": messages,
            "max_tokens": 25600,
            "temperature": 0.5,
            "n": 1,
            "stream": True
        }
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {OPENAI_API_KEY}"
        }
        full_response = ""
        with requests.post(OPENAI_BASE_URL + "/chat/completions", headers=headers, json=data, stream=True) as resp:
            for line in resp.iter_lines():
                if line:
                    if line.startswith(b'data: '):
                        line = line[6:]
                    if line == b'[DONE]':
                        break
                    try:
                        chunk = json.loads(line)
                        content = chunk.get('choices', [{}])[0].get('delta', {}).get('content', '')
                        if content:
                            full_response += content
                            yield content.encode('utf-8')
                    except Exception:
                        continue
        # After streaming, check for RFP proposal tag and generate PDF if present
        match = re.search(r'<RFP_PROPOSAL>(.*?)</RFP_PROPOSAL>', full_response, re.DOTALL)
        if match:
            proposal = match.group(1).strip()
            pdf_dir = os.path.join(os.path.dirname(__file__), 'static')
            if not os.path.exists(pdf_dir):
                os.makedirs(pdf_dir)
            pdf_filename = f"rfp_update_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            pdf_path = os.path.join(pdf_dir, pdf_filename)
            generate_pdf_from_text(proposal, pdf_path)
            download_url = f"/static/{pdf_filename}"
            msg = f"<h2>Updated Proposal PDF</h2><p>Your updated proposal has been generated as a PDF.</p><a href=\"{download_url}\" target=\"_blank\"><strong>Download Updated Proposal PDF</strong></a>"
            yield msg.encode('utf-8')
    return StreamingResponse(stream_response(), media_type="text/plain")

@app.get("/")
def serve_frontend():
    return FileResponse("simple_chat.html")

@app.get("/openapi.json", include_in_schema=False)
def custom_openapi():
    return get_openapi(
        title="BanglaChatbot API",
        version="1.0.0",
        description="API for BanglaChatbot with PDF context and streaming responses using Open LLM.",
        routes=app.routes,
    )

# Add a class to store chat context with embedded files
class ChatContext:
    def __init__(self):
        self.previous_questions = []
        self.previous_embeddings = []
        self.file_contexts = {}  # Map of file_id to file content context
    
    def add_question(self, question, embedding):
        self.previous_questions.append(question)
        self.previous_embeddings.append(embedding)
        # Keep only the last 10 questions for context
        if len(self.previous_questions) > 10:
            self.previous_questions.pop(0)
            self.previous_embeddings.pop(0)
    
    def add_file_context(self, file_id, context):
        self.file_contexts[file_id] = context
    
    def get_similar_questions(self, query_embedding, top_k=3):
        if not self.previous_embeddings:
            return []
        
        # Calculate cosine similarity
        similarities = []
        for i, emb in enumerate(self.previous_embeddings):
            similarity = np.dot(query_embedding, emb) / (np.linalg.norm(query_embedding) * np.linalg.norm(emb))
            similarities.append((i, similarity))
        
        # Sort by similarity (descending)
        similarities.sort(key=lambda x: x[1], reverse=True)
        
        # Return top_k similar questions
        return [(self.previous_questions[i], sim) for i, sim in similarities[:top_k] if sim > 0.7]
    
    def get_file_contexts(self):
        return list(self.file_contexts.values())

# Initialize chat context
chat_context = ChatContext()

# Modify the upload endpoint to store file context
@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    file_path = os.path.join(UPLOADS_PATH, file.filename)
    try:
        if not os.path.exists(UPLOADS_PATH):
            os.makedirs(UPLOADS_PATH)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext in [".xlsx", ".xls", ".csv"]:
            # --- Excel/Framework Agreement Processing ---
            if file_ext == ".csv":
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            # Create a summary (first 3 rows as string)
            summary = df.head(3).to_string(index=False)
            file_id = f"framework_{hash(file.filename)}_{len(chat_context.file_contexts)}"
            chat_context.add_file_context(file_id, {
                "filename": file.filename,
                "summary": summary,
                "upload_time": str(datetime.now())
            })
            # Embed each row as a string
            rows = df.astype(str).fillna("").values.tolist()
            row_texts = [" | ".join(row) for row in rows]
            row_embeddings = [model.encode(text) for text in row_texts]
            row_payloads = [{"row": text, "source": file.filename} for text in row_texts]
            start_id = qdrant_client.count(collection_name=FRAMEWORK_AGREEMENTS_COLLECTION).count
            qdrant_client.upsert(
                collection_name=FRAMEWORK_AGREEMENTS_COLLECTION,
                points=[
                    {
                        "id": start_id + i,
                        "vector": row_embeddings[i],
                        "payload": row_payloads[i],
                    }
                    for i in range(len(row_embeddings))
                ],
            )
            return JSONResponse(content={
                "filename": file.filename,
                "status": "success",
                "rows": len(row_texts),
                "message": f"Framework agreement Excel successfully uploaded and processed with {len(row_texts)} rows."
            })
        # --- End of Excel/Framework Agreement Processing ---

        # --- PDF Processing ---
        new_loader = PyPDFLoader(file_path)
        new_documents = new_loader.load()
        new_chunks = splitter.split_documents(new_documents)
        
        # Generate embeddings for new chunks
        new_embeddings = [model.encode(chunk.page_content) for chunk in new_chunks]
        new_payloads = [{"text": chunk.page_content, "source": file.filename} for chunk in new_chunks]
        
        # Store file context in chat context
        file_id = f"file_{hash(file.filename)}_{len(chat_context.file_contexts)}"
        file_summary = "\n".join([chunk.page_content for chunk in new_chunks[:3]])  # Use first 3 chunks as summary
        chat_context.add_file_context(file_id, {
            "filename": file.filename,
            "summary": file_summary,
            "upload_time": str(datetime.now())
        })
        
        # Add to the uploads collection in Qdrant instead of the main collection
        start_id = qdrant_client.count(collection_name=UPLOADS_COLLECTION_NAME).count
        qdrant_client.upsert(
            collection_name=UPLOADS_COLLECTION_NAME,
            points=[
                {
                    "id": start_id + i,
                    "vector": new_embeddings[i],
                    "payload": new_payloads[i],
                }
                for i in range(len(new_embeddings))
            ],
        )
        
        return JSONResponse(content={
            "filename": file.filename,
            "status": "success",
            "chunks": len(new_chunks),
            "message": f"Document successfully uploaded and processed into {len(new_chunks)} chunks."
        })
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "filename": file.filename,
                "status": "error",
                "message": f"Failed to process document: {str(e)}"
            }
        )

# Add endpoint to list available documents
@app.get("/documents")
async def list_documents():
    if not os.path.exists(DOCS_PATH):
        return {"documents": []}
        
    files = [f for f in os.listdir(DOCS_PATH) if f.endswith('.pdf')]
    return {"documents": files}

# Add endpoint to clear all vectors from the collection
@app.post("/clear-vectors")
async def clear_vectors():
    try:
        # Delete and recreate the collection
        try:
            qdrant_client.delete_collection(collection_name=ORIGINAL_COLLECTION_NAME)
        except Exception as e:
            print(f"Original collection delete skipped or failed: {e}")
            
        try:
            qdrant_client.create_collection(
                collection_name=ORIGINAL_COLLECTION_NAME,
                vectors_config=VectorParams(size=384, distance=Distance.COSINE),
            )
        except Exception as e:
            print(f"Original collection creation skipped or failed: {e}")
        
        # Reset the chat context as well
        global chat_context
        chat_context = ChatContext()
        
        return JSONResponse(content={
            "status": "success",
            "message": "Successfully cleared all vectors and chat context."
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "status": "error",
                "message": f"Failed to clear vectors: {str(e)}"
            }
        )

@app.post("/clear-all")
async def clear_all():
    try:
        # Delete all collections
        for collection in [ORIGINAL_COLLECTION_NAME, UPLOADS_COLLECTION_NAME, FRAMEWORK_AGREEMENTS_COLLECTION]:
            try:
                qdrant_client.delete_collection(collection_name=collection)
            except Exception as e:
                print(f"Collection {collection} delete skipped or failed: {e}")
            try:
                qdrant_client.create_collection(
                    collection_name=collection,
                    vectors_config=VectorParams(size=384, distance=Distance.COSINE),
                )
            except Exception as e:
                print(f"Collection {collection} creation skipped or failed: {e}")
        # Reset chat context
        global chat_context
        chat_context = ChatContext()
        return JSONResponse(content={
            "status": "success",
            "message": "Successfully cleared all vectors and chat context."
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "status": "error",
                "message": f"Failed to clear all vectors: {str(e)}"
            }
        )

from fastapi.staticfiles import StaticFiles
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.post("/rfp-chat")
def rfp_chat_endpoint(request: ChatRequest):
    def extract_text_from_pdf(pdf_path):
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(pdf_path)
            text = "\n".join(page.extract_text() or '' for page in reader.pages)
            return text
        except Exception as e:
            return f"[Error extracting PDF text: {e}]"

    def get_guideline_context():
        guideline_query = "BRAC RFP guideline"
        guideline_vec = model.encode(guideline_query)
        guideline_hits = qdrant_client.search(
            collection_name=ORIGINAL_COLLECTION_NAME,
            query_vector=guideline_vec,
            limit=3,
        )
        return "\n".join([hit.payload['text'] for hit in guideline_hits])

    def stream_response():
        # Build chat history for LLM
        messages = [
            {"role": "system", "content": '''You are BRACGPT, an expert RFP (Request for Proposal) builder and verifier for BRAC. Your job is to help the user build, review, and verify RFP documents step by step.\n\nINSTRUCTIONS:\n1. Guide the user through the RFP creation process, one section at a time (e.g., Title, Background, Requirements, Evaluation Criteria, Submission Instructions, etc.).\n2. If the user uploads an RFP, analyze and verify it against BRAC's RFP guidelines.\n3. Always use the full chat history for context.\n4. After gathering all necessary details, generate a full RFP proposal in clear, structured English, using HTML formatting for sections and lists.\n5. Enclose the final RFP proposal in <RFP_PROPOSAL>...</RFP_PROPOSAL> tags.\n6. If verifying, provide a detailed report of compliance and suggestions.\n7. Do not answer general procurement questions or provide unrelated advice. Only focus on building and verifying the RFP.\n8. Be concise, professional, and use BRAC's organizational tone.\n9. If the user says they are done or ready, proceed to generate or verify the full RFP.\n10. Always provide the full proposal inside <RFP_PROPOSAL>...</RFP_PROPOSAL> tags when generating a proposal.'''}
        ]
        # Add chat history
        for q, a in request.chat_history:
            messages.append({"role": "user", "content": q})
            messages.append({"role": "assistant", "content": a})
        # Add current user question
        messages.append({"role": "user", "content": request.question})

        # If user wants to verify an uploaded RFP
        if request.rfp_info and request.rfp_info.get("uploaded_rfp_path"):
            rfp_path = request.rfp_info["uploaded_rfp_path"]
            rfp_text = extract_text_from_pdf(rfp_path)
            guideline_text = get_guideline_context()
            messages.append({
                "role": "user",
                "content": f"Please verify the following RFP against BRAC's RFP guidelines.\n\nRFP Text:\n{rfp_text}\n\nGuidelines:\n{guideline_text}"
            })

        data = {
            "model": OPENAI_MODEL,
            "messages": messages,
            "max_tokens": 25600,
            "temperature": 0.4,
            "n": 1,
            "stream": True
        }
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {OPENAI_API_KEY}"
        }
        full_response = ""
        with requests.post(OPENAI_BASE_URL + "/chat/completions", headers=headers, json=data, stream=True) as resp:
            for line in resp.iter_lines():
                if line:
                    if line.startswith(b'data: '):
                        line = line[6:]
                    if line == b'[DONE]':
                        break
                    try:
                        chunk = json.loads(line)
                        content = chunk.get('choices', [{}])[0].get('delta', {}).get('content', '')
                        if content:
                            full_response += content
                            yield content.encode('utf-8')
                    except Exception:
                        continue
        # After streaming, check for RFP proposal tag and generate PDF if present
        if request.question.strip().lower() == "download":
            # Find the most recent RFP PDF or Markdown file in the static directory
            pdf_dir = os.path.join(os.path.dirname(__file__), 'static')
            files = [f for f in os.listdir(pdf_dir) if (f.startswith('rfp_') and (f.endswith('.pdf') or f.endswith('.md')))]
            if files:
                latest_file = max(files, key=lambda f: os.path.getctime(os.path.join(pdf_dir, f)))
                download_url = f"/static/{latest_file}"
                if latest_file.endswith('.md'):
                    msg = f"<h2>RFP Proposal (Markdown)</h2><p>Your RFP proposal is ready as a Markdown file.</p><a href=\"{download_url}\" target=\"_blank\"><strong>Download RFP Markdown</strong></a><br><br><iframe src=\"{download_url}\" width=\"100%\" height=\"600\" style=\"border:1px solid #ccc; background:#fafbfc;\"></iframe>"
                else:
                    msg = f"<h2>RFP Proposal PDF</h2><p>Your RFP proposal is ready.</p><a href=\"{download_url}\" target=\"_blank\"><strong>Download RFP PDF</strong></a>"
                yield msg.encode('utf-8')
                return
        match = re.search(r'<RFP_PROPOSAL>(.*?)</RFP_PROPOSAL>', full_response, re.DOTALL)
        if match:
            proposal = match.group(1).strip()
            pdf_dir = os.path.join(os.path.dirname(__file__), 'static')
            if not os.path.exists(pdf_dir):
                os.makedirs(pdf_dir)
            pdf_filename = f"rfp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            pdf_path = os.path.join(pdf_dir, pdf_filename)
            # Convert Markdown to HTML, then HTML to PDF
            import markdown2
            from weasyprint import HTML
            html_content = markdown2.markdown(proposal)
            HTML(string=html_content).write_pdf(pdf_path)
            download_url = f"/static/{pdf_filename}"
            msg = f"<h2>RFP Proposal PDF</h2><p>Your RFP proposal has been generated as a PDF.</p><a href=\"{download_url}\" target=\"_blank\"><strong>Download RFP PDF</strong></a>"
            yield msg.encode('utf-8')
    return StreamingResponse(stream_response(), media_type="text/plain")